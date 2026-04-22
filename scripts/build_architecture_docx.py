"""One-off: build docs/ARCHITECTURE_AND_CODE_MAP.docx from docs/ARCHITECTURE_AND_CODE_MAP.md then delete the .md."""
from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "ARCHITECTURE_AND_CODE_MAP.md"
DOCX_PATH = ROOT / "docs" / "ARCHITECTURE_AND_CODE_MAP.docx"


def add_para_with_bold(doc, text: str, style: str | None = None) -> None:
    """Split on **bold** markers and add one paragraph with mixed runs."""
    text = text.replace("\t", " ").strip()
    if not text:
        return
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def parse_table_row(line: str) -> list[str]:
    inner = line.strip()[1:-1]
    return [c.strip() for c in inner.split("|")]


def is_separator_row(line: str) -> bool:
    cells = parse_table_row(line)
    return all(re.fullmatch(r":?-{3,}:?", c or "-") for c in cells)


def main() -> int:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        print("Install python-docx: pip install python-docx", file=sys.stderr)
        return 1

    if not MD_PATH.is_file():
        print(f"Missing source: {MD_PATH}", file=sys.stderr)
        return 1

    raw = MD_PATH.read_text(encoding="utf-8")
    lines = raw.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(12)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if is_table_row(line):
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row = parse_table_row(lines[i])
                if not is_separator_row(lines[i]):
                    rows.append(row)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cell_text = row[ci] if ci < len(row) else ""
                        # Table cells: plain text (strip markdown bold markers)
                        table.rows[ri].cells[ci].text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                doc.add_paragraph()
            continue

        if line.strip() == "":
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(line.strip(), style="List Number")
            for r in p.runs:
                r.font.name = "Calibri"
            i += 1
            continue

        if line.strip().startswith("- "):
            add_para_with_bold(doc, line.strip()[2:], style="List Bullet")
            i += 1
            continue

        st = line.strip()
        if (
            st.startswith("*")
            and st.endswith("*")
            and not st.startswith("**")
            and len(st) >= 2
        ):
            p = doc.add_paragraph()
            r = p.add_run(st[1:-1])
            r.italic = True
            i += 1
            continue

        add_para_with_bold(doc, line.strip())
        i += 1

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")
    MD_PATH.unlink()
    print(f"Removed {MD_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
